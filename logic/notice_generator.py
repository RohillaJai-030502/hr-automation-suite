# ============================================================
# notice_generator.py — Generates DOCX Notice (Single Page)
# ============================================================

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import math

OUTPUT_DIR = "generated_notices"

def set_paragraph_spacing(para, before=0, after=0, line_spacing=None):
    """Set tight spacing for paragraphs"""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top","left","bottom","right","insideH","insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_centered_bold(doc, text, size=11):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    set_paragraph_spacing(para, before=0, after=0, line_spacing=11)
    return para

def generate_notice_docx(data):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()

    # ── Tight Page Margins ──
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── HEADER ──
    h1 = add_centered_bold(doc, "अन्तर कार्यालय नोट / ION", size=12)
    set_paragraph_spacing(h1, before=0, after=4)

    h2 = add_centered_bold(doc, "चरम प्राक्षेपिकी अनुसंधान प्रयोगशाला / TBRL", size=11)
    set_paragraph_spacing(h2, before=0, after=4)

    h3 = add_centered_bold(doc, "{मानव संसाधन विकास विभाग / HRD}", size=11)
    set_paragraph_spacing(h3, before=0, after=4)

    # Small spacer
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=2, after=2, line_spacing=8)

    # ── SUBJECT ──
    subject_para = doc.add_paragraph()
    subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject_run = subject_para.add_run(
        f"विषय : {data['degree']} छात्रों की औद्योगिक प्रशिक्षण"
    )
    subject_run.bold = True
    subject_run.font.size = Pt(11)
    subject_run.underline = True
    set_paragraph_spacing(subject_para, before=2, after=4)

    # ── HINDI BODY ──
    hindi_body = (
        f"{data['start_month']} {data['start_year']} से {data['end_month']} "
        f"{data['end_year']} तक प्रशिक्षण प्राप्त करने के लिये विभिन्न विश्वविद्यालय से "
        f"{data['degree']} के छात्र टी.बी.आर.एल में प्रशिक्षण प्राप्त करने के लिये आवेदन करेंगे।\n"
        f"जो वैज्ञानिक उन्हे प्रशिक्षण देना चाहते हैं, वे अपनी आवश्यकता मानव संसाधन विकास विभाग को "
        f"{data['last_date']} तक भेजे । अपनी आवश्यकता को वेAD/TD/GD/JD/PD या वरिष्ठतम "
        f"वैज्ञानिक से अग्रेसित अवश्य करायें ।"
    )
    hindi_para = doc.add_paragraph(hindi_body)
    hindi_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    hindi_para.runs[0].font.size = Pt(10)
    set_paragraph_spacing(hindi_para, before=2, after=4)

    # ── ENGLISH BODY 1 ──
    english_body1 = (
        f"{data['degree']} Students from different University interested in getting training at TBRL "
        f"during {data['start_month']}-{data['end_month']} {data['end_year']} will submit their applications."
    )
    eng_para1 = doc.add_paragraph(english_body1)
    eng_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    eng_para1.runs[0].font.size = Pt(10)
    set_paragraph_spacing(eng_para1, before=2, after=4)

    # ── ENGLISH BODY 2 ──
    english_body2 = (
        f"Scientists who want to impart training to these students are requested to give their "
        f"requirement to HRD Division by {data['last_date']}, forwarded by AD/TD/GD/JD/PD or "
        f"senior most scientist of the group."
    )
    eng_para2 = doc.add_paragraph(english_body2)
    eng_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    eng_para2.runs[0].font.size = Pt(10)
    set_paragraph_spacing(eng_para2, before=2, after=6)

    # ── SIGNATURE SPACE ──
    for _ in range(3):
        spacer = doc.add_paragraph()
        set_paragraph_spacing(spacer, before=0, after=0, line_spacing=8)

    # ── ION NUMBER + DATE + SIGNATORY ──
    info_table = doc.add_table(rows=2, cols=2)

    left_cell = info_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.add_run(f"ION नं0: {data['ion_number']}").font.size = Pt(10)
    set_paragraph_spacing(left_para, before=0, after=2)

    date_cell = info_table.cell(1, 0)
    date_para = date_cell.paragraphs[0]
    date_para.add_run(f"दिनांक:    {data['notice_date']}").font.size = Pt(10)
    set_paragraph_spacing(date_para, before=0, after=2)

    right_cell = info_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_para.add_run(f"({data['signatory_name']})")
    right_run.bold = True
    right_run.font.size = Pt(10)
    set_paragraph_spacing(right_para, before=0, after=2)

    desig_cell = info_table.cell(1, 1)
    desig_para = desig_cell.paragraphs[0]
    desig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    desig_para.add_run(data['signatory_designation']).font.size = Pt(10)
    set_paragraph_spacing(desig_para, before=0, after=2)

    # ── TO SECTION ──
    spacer2 = doc.add_paragraph()
    set_paragraph_spacing(spacer2, before=2, after=2, line_spacing=6)

    to_para = doc.add_paragraph()
    to_run = to_para.add_run("To:")
    to_run.bold = True
    to_run.font.size = Pt(10)
    set_paragraph_spacing(to_para, before=2, after=0)

    to_all = doc.add_paragraph("        All TD's / GD's / PD's")
    to_all.runs[0].font.size = Pt(10)
    set_paragraph_spacing(to_all, before=0, after=4)

    # ── DEPARTMENTS TABLE ──
    departments = data["departments"]
    if departments:
        num_cols = min(5, math.ceil(math.sqrt(len(departments) * 2)))
        num_rows = math.ceil(len(departments) / num_cols)

        while len(departments) < num_rows * num_cols:
            departments.append("")

        dept_table = doc.add_table(rows=num_rows, cols=num_cols)
        dept_table.style = "Table Grid"

        for i, dept in enumerate(departments):
            row = i // num_cols
            col = i % num_cols
            cell = dept_table.cell(row, col)
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell_para.add_run(dept)
            run.font.size = Pt(9)
            set_paragraph_spacing(cell_para, before=2, after=2)

    # ── NOTE ──
    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        "Note: Performa appended overleaf (the soft copy of the performa "
        "can be downloaded from TBRL portal)"
    )
    note_run.underline = True
    note_run.font.size = Pt(9)
    set_paragraph_spacing(note_para, before=4, after=0)

    # ── SAVE ──
    filename = f"ION_Notice_{data['start_month']}_{data['end_month']}_{data['end_year']}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath