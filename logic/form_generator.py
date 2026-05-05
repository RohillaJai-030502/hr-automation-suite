# ============================================================
# form_generator.py — Generates FM/HRD-09 Requirement Form
# ============================================================

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "generated_notices"

def set_cell_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def generate_form_docx():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()

    # ── Page Margins ──
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.5)
        section.right_margin  = Inches(0.5)

    # ── Space at top (4 lines for form name) ──
    for _ in range(4):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0, line=10)

    # ══════════════════════════════════════
    # 1. HEADER TABLE
    # ══════════════════════════════════════
    header_table = doc.add_table(rows=2, cols=2)
    header_table.style = 'Table Grid'

    # Row 1 — merged title
    row1 = header_table.rows[0]
    row1.cells[0].merge(row1.cells[1])
    title_cell = header_table.rows[0].cells[0]
    set_cell_height(row1, 0.8)
    title_para = title_cell.paragraphs[0]
    set_para_spacing(title_para, before=2, after=2)
    title_run = title_para.add_run('Requirement For Industrial Trainees')
    title_run.bold = True
    title_run.font.size = Pt(12)
    tab_run = title_para.add_run('\t(Separate for each Discipline/Branch)')
    tab_run.font.size = Pt(9)

    # Row 2 — FM/HRD-09 | Rev: 00
    row2 = header_table.rows[1]
    set_cell_height(row2, 0.6)
    set_col_width(row2.cells[0], 11.5)
    set_col_width(row2.cells[1], 6.9)

    p_left = row2.cells[0].paragraphs[0]
    set_para_spacing(p_left, before=2, after=2)
    r = p_left.add_run('FM/HRD-09')
    r.bold = True
    r.font.size = Pt(10)

    p_right = row2.cells[1].paragraphs[0]
    set_para_spacing(p_right, before=2, after=2)
    r2 = p_right.add_run('Rev: 00')
    r2.bold = True
    r2.font.size = Pt(10)

    spacer = doc.add_paragraph()
    set_para_spacing(spacer, before=0, after=4, line=6)

    # ══════════════════════════════════════
    # 2. GROUP & SESSION TABLE
    # ══════════════════════════════════════
    group_table = doc.add_table(rows=1, cols=2)
    group_table.style = 'Table Grid'

    grp_left  = group_table.rows[0].cells[0]
    grp_right = group_table.rows[0].cells[1]

    set_col_width(grp_left, 11.5)
    set_col_width(grp_right, 6.9)
    set_cell_height(group_table.rows[0], 3.2)

    grp_left_para = grp_left.paragraphs[0]
    set_para_spacing(grp_left_para, before=2, after=2)
    grp_left_run = grp_left_para.add_run('1. Group Name / Project:')
    grp_left_run.bold = True
    grp_left_run.font.size = Pt(10)

    grp_right_para = grp_right.paragraphs[0]
    set_para_spacing(grp_right_para, before=2, after=2)
    grp_right_run = grp_right_para.add_run('Training Session:')
    grp_right_run.bold = True
    grp_right_run.font.size = Pt(10)

    cb1 = grp_right.add_paragraph()
    set_para_spacing(cb1, before=8, after=4)
    cb1.add_run('☐  Jan - June').font.size = Pt(10)

    cb2 = grp_right.add_paragraph()
    set_para_spacing(cb2, before=8, after=4)
    cb2.add_run('☐  July - Dec').font.size = Pt(10)

    spacer2 = doc.add_paragraph()
    set_para_spacing(spacer2, before=0, after=4, line=6)

    # ══════════════════════════════════════
    # 3. DISCIPLINE TABLE
    # ══════════════════════════════════════
    disc_header = doc.add_paragraph()
    set_para_spacing(disc_header, before=0, after=4)
    disc_run = disc_header.add_run('2. Discipline and Requirements:')
    disc_run.bold = True
    disc_run.font.size = Pt(10)

    disciplines = [
        'Computer Science',
        'Mechanical',
        'Civil Engineering',
        'ECE',
        'Chemical / Chemistry',
        'Aerospace',
        'Others: _______________',
    ]

    disc_table = doc.add_table(rows=len(disciplines)+1, cols=4)
    disc_table.style = 'Table Grid'

    col_widths = [5.0, 1.8, 2.5, 9.1]

    # Header row
    headers = ['Discipline / Branch', 'Select', 'No. of Trainees', 'Work / Project to be Assigned']
    header_row = disc_table.rows[0]
    set_cell_height(header_row, 0.8)
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        set_col_width(cell, col_widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before=2, after=2)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Data rows — 1.4cm each to fill page
    for i, disc in enumerate(disciplines):
        row = disc_table.rows[i+1]
        set_cell_height(row, 1.4)

        cell0 = row.cells[0]
        set_col_width(cell0, col_widths[0])
        p0 = cell0.paragraphs[0]
        set_para_spacing(p0, before=4, after=2)
        p0.add_run(disc).font.size = Pt(10)

        cell1 = row.cells[1]
        set_col_width(cell1, col_widths[1])
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p1, before=4, after=2)
        p1.add_run('☐').font.size = Pt(12)

        cell2 = row.cells[2]
        set_col_width(cell2, col_widths[2])

        cell3 = row.cells[3]
        set_col_width(cell3, col_widths[3])

    spacer3 = doc.add_paragraph()
    set_para_spacing(spacer3, before=0, after=4, line=6)

    # ══════════════════════════════════════
    # 4. SKILLS & INSTITUTE TABLE
    # ══════════════════════════════════════
    skills_table = doc.add_table(rows=1, cols=2)
    skills_table.style = 'Table Grid'

    skills_left  = skills_table.rows[0].cells[0]
    skills_right = skills_table.rows[0].cells[1]
    set_cell_height(skills_table.rows[0], 5.0)

    p_sl = skills_left.paragraphs[0]
    set_para_spacing(p_sl, before=2, after=2)
    r_sl = p_sl.add_run('3. Desired Skills (If any):')
    r_sl.bold = True
    r_sl.font.size = Pt(10)

    p_sr = skills_right.paragraphs[0]
    set_para_spacing(p_sr, before=2, after=2)
    r_sr = p_sr.add_run('4. Recommended Institute (If any):')
    r_sr.bold = True
    r_sr.font.size = Pt(10)

    spacer4 = doc.add_paragraph()
    set_para_spacing(spacer4, before=0, after=6, line=8)

    # ══════════════════════════════════════
    # 5. SIGNATURE FOOTER
    # ══════════════════════════════════════
    sig_table = doc.add_table(rows=3, cols=3)
    remove_table_borders(sig_table)

    # Row 1
    r1 = sig_table.rows[0]
    set_cell_height(r1, 0.7)
    p_date = r1.cells[0].paragraphs[0]
    set_para_spacing(p_date, before=2, after=2)
    date_run = p_date.add_run('Date: ________________________')
    date_run.bold = True
    date_run.font.size = Pt(10)

    p_sig = r1.cells[2].paragraphs[0]
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p_sig, before=2, after=2)
    p_sig.add_run('______________________________').font.size = Pt(10)

    # Row 2
    r2 = sig_table.rows[1]
    set_cell_height(r2, 0.6)
    p_name = r2.cells[2].paragraphs[0]
    p_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p_name, before=2, after=2)
    name_run = p_name.add_run('Name & Signature')
    name_run.bold = True
    name_run.font.size = Pt(10)

    # Row 3
    r3 = sig_table.rows[2]
    set_cell_height(r3, 0.6)
    p_desig = r3.cells[2].paragraphs[0]
    p_desig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p_desig, before=2, after=2)
    desig_run = p_desig.add_run('Designation: ____________________')
    desig_run.bold = True
    desig_run.font.size = Pt(10)

    # ── Save ──
    filename = "FM_HRD_09_Requirement_Form.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath